import streamlit as st
import pandas as pd
import base64
import random
import json
from datetime import datetime
import io
import os
import tempfile
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfkit

# Page configuration
st.set_page_config(
    page_title="🎨 AI Resume Generator - 1000+ Templates",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 15px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
        box-shadow: 0 10px 30px rgba(0,0,0,0.2);
    }
    
    .template-card {
        background: white;
        border-radius: 15px;
        padding: 1.2rem;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        transition: all 0.3s;
        cursor: pointer;
        border: 3px solid transparent;
        margin-bottom: 15px;
    }
    
    .template-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 15px 30px rgba(0,0,0,0.15);
    }
    
    .template-card.selected {
        border-color: #667eea;
        background: #f8f9ff;
    }
    
    .template-preview {
        height: 100px;
        border-radius: 8px;
        margin-bottom: 10px;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 24px;
        color: white;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
    }
    
    .download-btn {
        background: #667eea;
        color: white;
        padding: 10px;
        border-radius: 5px;
        border: none;
        width: 100%;
        cursor: pointer;
        font-weight: 600;
        margin: 2px 0;
    }
    
    .download-link {
        display: block;
        background: #28a745;
        color: white;
        padding: 10px;
        border-radius: 5px;
        text-align: center;
        text-decoration: none;
        margin: 5px 0;
        font-weight: 600;
    }
    
    .download-link:hover {
        background: #218838;
    }
    
    .footer {
        text-align: center;
        padding: 2rem;
        color: #666;
        margin-top: 3rem;
        border-top: 1px solid #e0e0e0;
    }
    
    .delete-btn {
        background: #ff4757;
        color: white;
        border: none;
        border-radius: 50%;
        width: 30px;
        height: 30px;
        display: flex;
        align-items: center;
        justify-content: center;
        cursor: pointer;
    }
    
    .add-template-btn {
        background: #28a745;
        color: white;
        padding: 10px;
        border-radius: 5px;
        border: none;
        width: 100%;
        cursor: pointer;
        font-weight: 600;
        margin: 10px 0;
    }
    
    .stTabs [data-baseweb="tab-list"] {
        gap: 2rem;
    }
    
    .stTabs [data-baseweb="tab"] {
        height: 3rem;
        white-space: pre-wrap;
        background-color: #f0f2f6;
        border-radius: 5px;
        padding: 0.5rem 1rem;
    }
    
    .education-field {
        border: 1px solid #e0e0e0;
        padding: 1rem;
        border-radius: 8px;
        margin-bottom: 1rem;
    }
</style>
""", unsafe_allow_html=True)

# Helper function for safe dictionary access
def safe_get(data, *keys, default=''):
    """Safely get nested dictionary values without throwing errors"""
    try:
        for key in keys:
            if isinstance(data, dict):
                data = data.get(key, default)
            else:
                return default
        return data if data is not None else default
    except (TypeError, AttributeError):
        return default

# Initialize session state with BLANK fields
def init_session_state():
    """Initialize or reset session state with BLANK data"""
    default_data = {
        'personal': {
            'name': '',
            'email': '',
            'phone': '',
            'address': '',
            'city': '',
            'state': '',
            'pin': '',
            'linkedin': '',
            'website': '',
            'title': ''
        },
        'profile_image': None,
        'summary': '',
        'objective': '',
        'experience': [],
        'education': [],  # Will store flexible education entries
        'skills': {},
        'certifications': [],
        'languages': [],
        'achievements': [],
        'personal_details': {
            'father_name': '',
            'dob': '',
            'marital_status': '',
            'nationality': '',
            'gender': ''
        },
        'custom_sections': []
    }
    
    st.session_state.resume_data = default_data
    
    if 'selected_template' not in st.session_state:
        st.session_state.selected_template = 'style1_v1'
    
    if 'template_category' not in st.session_state:
        st.session_state.template_category = 'All'
    
    if 'photo' not in st.session_state:
        st.session_state.photo = None
    
    if 'template_family' not in st.session_state:
        st.session_state.template_family = 'Style 1: Minimalist Luxury'
    
    if 'pdf_data' not in st.session_state:
        st.session_state.pdf_data = None
    
    if 'word_data' not in st.session_state:
        st.session_state.word_data = None
    
    if 'custom_templates' not in st.session_state:
        st.session_state.custom_templates = {}
    
    # Education field configuration - defines what fields each education entry has
    if 'edu_fields' not in st.session_state:
        st.session_state.edu_fields = ['degree', 'institution', 'year', 'grade', 'specialization']

# Safe reset function
def safe_reset():
    """Safely reset all session state to blank"""
    # Clear all existing session state
    for key in list(st.session_state.keys()):
        del st.session_state[key]
    # Reinitialize with blank data
    init_session_state()

# Generate 105+ Template Variations from the 3 Base Styles
def generate_template_variations():
    templates = {}
    
    # Base Style 1: Minimalist Luxury
    style1_base = {
        'name': 'Minimalist Luxury',
        'family': 'Style 1: Minimalist Luxury',
        'primary_colors': ['#1e2b3a', '#2c3e50', '#34495e', '#2d3e4f', '#1a2b3c'],
        'secondary_colors': ['#ff6b6b', '#ff8787', '#fa5252', '#ff5f5f', '#ff4d4d'],
        'accent_colors': ['#ffd93d', '#fcc419', '#ffd43b', '#f59f00', '#fab005']
    }
    
    # Base Style 2: Modern Corporate
    style2_base = {
        'name': 'Modern Corporate',
        'family': 'Style 2: Modern Corporate',
        'primary_colors': ['#0b1e33', '#1e3a5f', '#0f2b45', '#1a3a5a', '#0d2b40'],
        'secondary_colors': ['#00c3ff', '#339af0', '#228be6', '#4dabf7', '#3b9eff'],
        'accent_colors': ['#e0f0ff', '#d0ebff', '#a5d8ff', '#74c0fc', '#4dabf7']
    }
    
    # Base Style 3: Creative Hospitality
    style3_base = {
        'name': 'Creative Hospitality',
        'family': 'Style 3: Creative Hospitality',
        'primary_colors': ['#c49a6c', '#a67b5b', '#b68b6b', '#8b6b4b', '#9b7b5b'],
        'secondary_colors': ['#5a3e2b', '#6b4e3b', '#7b5e4b', '#4a3a2a', '#3a2a1a'],
        'accent_colors': ['#fff0e6', '#ffe4cc', '#ffd9b3', '#ffcc99', '#ffbf80']
    }
    
    # Generate 35 variations for Style 1
    for i in range(35):
        templates[f"style1_v{i+1}"] = {
            'id': f"style1_v{i+1}",
            'name': f"Minimalist Luxury {i+1}",
            'family': 'Built-in',
            'style_family': 'Style 1: Minimalist Luxury',
            'colors': {
                'primary': random.choice(style1_base['primary_colors']),
                'secondary': random.choice(style1_base['secondary_colors']),
                'accent': random.choice(style1_base['accent_colors'])
            },
            'is_custom': False
        }
    
    # Generate 35 variations for Style 2
    for i in range(35):
        templates[f"style2_v{i+1}"] = {
            'id': f"style2_v{i+1}",
            'name': f"Modern Corporate {i+1}",
            'family': 'Built-in',
            'style_family': 'Style 2: Modern Corporate',
            'colors': {
                'primary': random.choice(style2_base['primary_colors']),
                'secondary': random.choice(style2_base['secondary_colors']),
                'accent': random.choice(style2_base['accent_colors'])
            },
            'is_custom': False
        }
    
    # Generate 35 variations for Style 3
    for i in range(35):
        templates[f"style3_v{i+1}"] = {
            'id': f"style3_v{i+1}",
            'name': f"Creative Hospitality {i+1}",
            'family': 'Built-in',
            'style_family': 'Style 3: Creative Hospitality',
            'colors': {
                'primary': random.choice(style3_base['primary_colors']),
                'secondary': random.choice(style3_base['secondary_colors']),
                'accent': random.choice(style3_base['accent_colors'])
            },
            'is_custom': False
        }
    
    return templates

# Initialize templates
TEMPLATES = generate_template_variations()

# Function to add custom template
def add_custom_template(name, primary_color, secondary_color, accent_color, style_type):
    """Add a new custom template to the collection"""
    template_id = f"custom_{datetime.now().strftime('%Y%m%d%H%M%S')}"
    
    # Map style type to style family
    style_families = {
        'Minimalist Luxury': 'Style 1: Minimalist Luxury',
        'Modern Corporate': 'Style 2: Modern Corporate',
        'Creative Hospitality': 'Style 3: Creative Hospitality',
        'Custom Mix': 'Custom Styles'
    }
    
    family = style_families.get(style_type, 'Custom Styles')
    
    new_template = {
        'id': template_id,
        'name': name,
        'family': 'Custom',
        'style_family': family,
        'colors': {
            'primary': primary_color,
            'secondary': secondary_color,
            'accent': accent_color
        },
        'is_custom': True
    }
    
    # Add to session state custom templates
    if 'custom_templates' not in st.session_state:
        st.session_state.custom_templates = {}
    
    st.session_state.custom_templates[template_id] = new_template
    
    return template_id

# Function to get all templates (built-in + custom)
def get_all_templates():
    """Combine built-in and custom templates"""
    all_templates = TEMPLATES.copy()
    if 'custom_templates' in st.session_state:
        all_templates.update(st.session_state.custom_templates)
    return all_templates

def generate_id(prefix):
    return f"{prefix}_{datetime.now().strftime('%Y%m%d%H%M%S%f')}"

# ========== PDF-FRIENDLY STYLE 1: MINIMALIST LUXURY ==========
def generate_style1_html(template_id, data, template_colors=None):
    if template_colors:
        colors = template_colors
    else:
        all_templates = get_all_templates()
        template = all_templates.get(template_id)
        colors = template['colors']
    
    years_exp = 13
    total_skills = sum(len(skills) for skills in data['skills'].values()) if data['skills'] else 0
    
    # Photo handling
    photo_html = ""
    if data.get('profile_image'):
        photo_html = f'<img src="{data["profile_image"]}" style="width:100%; height:100%; border-radius:50%; object-fit:cover;" alt="Profile">'
    else:
        photo_html = f'<span style="font-size:60px; color:white;">{data["personal"]["name"][0] if data["personal"]["name"] else "?"}</span>'
    
    # Build skills as comma-separated text
    all_skills = []
    if data['skills']:
        for category, skills in data['skills'].items():
            all_skills.extend(skills)
    skills_text = " • ".join(all_skills) if all_skills else "Add skills in the editor"
    
    # Build education - Flexible fields
    edu_html = ""
    for edu in data['education']:
        edu_html += f"""
            <div style="margin-bottom: 15px; padding-left: 10px; border-left: 3px solid {colors['secondary']};">"""
        
        # Display all education fields dynamically
        if edu.get('year'):
            edu_html += f'<div style="color: {colors["accent"]}; font-weight: bold;">{edu["year"]}</div>'
        if edu.get('degree'):
            edu_html += f'<div style="font-weight: bold;">{edu["degree"]}</div>'
        if edu.get('institution'):
            edu_html += f'<div style="color: #888;">{edu["institution"]}</div>'
        if edu.get('grade'):
            edu_html += f'<div style="color: {colors["accent"]};">{edu["grade"]}</div>'
        if edu.get('specialization'):
            edu_html += f'<div style="color: #666; font-size: 13px;">{edu["specialization"]}</div>'
        
        edu_html += "</div>"
    
    if not data['education']:
        edu_html = '<div style="color: #888; font-style: italic;">No education added. Click "Edit All Sections" to add.</div>'
    
    # Build certifications
    cert_html = ""
    for cert in data['certifications']:
        cert_html += f"""
            <div style="margin-bottom: 15px; padding-left: 10px; border-left: 3px solid {colors['secondary']};">
                <div style="color: {colors['accent']}; font-weight: bold;">{cert['year'] if cert.get('year') else 'Year'}</div>
                <div style="font-weight: bold;">{cert['name'] if cert.get('name') else 'Certification'}</div>
                <div style="color: #888;">{cert['issuer'] if cert.get('issuer') else 'Issuer'}</div>
            </div>
        """
    if not data['certifications']:
        cert_html = '<div style="color: #888; font-style: italic;">No certifications added.</div>'
    
    # Build languages
    lang_text = "  •  ".join([f"{l['name']} ({l['proficiency']})" for l in data['languages']]) if data['languages'] else "Add languages in editor"
    
    # Build experience
    exp_html = ""
    for exp in data['experience']:
        exp_html += f"""
            <div style="margin-bottom: 25px;">
                <div style="display: flex; justify-content: space-between; margin-bottom: 5px; flex-wrap: wrap;">
                    <span style="font-size: 18px; font-weight: bold; color: {colors['primary']};">{exp['company'] if exp.get('company') else 'Company'}</span>
                    <span style="color: {colors['secondary']}; font-weight: 600;">{exp['start_date'] if exp.get('start_date') else 'Start'} - {exp['end_date'] if exp.get('end_date') else 'End'}</span>
                </div>
                <div style="font-size: 16px; color: #666; margin-bottom: 10px; font-style: italic;">{exp['position'] if exp.get('position') else 'Position'} | {exp['location'] if exp.get('location') else 'Location'}</div>
                <ul style="margin-left: 20px; margin-top: 5px;">
        """
        if exp.get('description'):
            for desc in exp['description']:
                exp_html += f"<li style='margin-bottom: 5px;'>{desc}</li>"
        else:
            exp_html += "<li style='color: #888;'>Add description</li>"
        exp_html += "</ul></div>"
    if not data['experience']:
        exp_html = '<div style="color: #888; font-style: italic;">No experience added. Click "Edit All Sections" to add.</div>'
    
    # Build achievements
    ach_html = ""
    for achievement in data['achievements']:
        ach_html += f"<li style='margin-bottom: 5px;'>{achievement}</li>"
    if not data['achievements']:
        ach_html = '<li style="color: #888; font-style: italic;">No achievements added.</li>'
    
    personal = data['personal_details']
    
    html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{data['personal']['name'] if data['personal']['name'] else 'Your Name'} - Resume</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        body {{
            font-family: Arial, Helvetica, sans-serif;
            background: #f0f0f0;
            padding: 30px;
        }}
        .resume {{
            max-width: 1100px;
            margin: 0 auto;
            background: white;
            box-shadow: 0 10px 30px rgba(0,0,0,0.2);
            display: table;
            width: 100%;
        }}
        .left-panel {{
            display: table-cell;
            width: 35%;
            background: {colors['primary']};
            color: white;
            padding: 30px;
            vertical-align: top;
        }}
        .right-panel {{
            display: table-cell;
            width: 65%;
            background: white;
            padding: 30px;
            vertical-align: top;
        }}
        .profile-image {{
            width: 150px;
            height: 150px;
            background: {colors['secondary']};
            border-radius: 50%;
            margin: 0 auto 20px;
            display: flex;
            align-items: center;
            justify-content: center;
            overflow: hidden;
            border: 4px solid {colors['accent']};
        }}
        .section-title {{
            font-size: 20px;
            font-weight: bold;
            margin: 25px 0 15px;
            border-bottom: 3px solid {colors['accent']};
            padding-bottom: 5px;
            color: {colors['primary']};
        }}
        .left-panel .section-title {{
            color: white;
            border-bottom-color: {colors['accent']};
        }}
        .contact-item {{
            margin-bottom: 10px;
            font-size: 14px;
        }}
        .left-panel .contact-item {{
            color: #f0f0f0;
        }}
        .stats-grid {{
            display: flex;
            flex-wrap: wrap;
            gap: 15px;
            margin: 20px 0;
        }}
        .stat-item {{
            flex: 1 1 calc(50% - 15px);
            background: #f8f9fa;
            padding: 15px;
            text-align: center;
            border-bottom: 3px solid {colors['accent']};
        }}
        .stat-number {{
            font-size: 24px;
            font-weight: bold;
            color: {colors['secondary']};
        }}
        .stat-label {{
            font-size: 12px;
            color: #666;
        }}
        .name {{
            font-size: 42px;
            font-weight: bold;
            color: {colors['primary']};
            margin-bottom: 5px;
        }}
        .title {{
            font-size: 18px;
            color: {colors['secondary']};
            margin-bottom: 20px;
            border-bottom: 2px solid #f0f0f0;
            padding-bottom: 15px;
        }}
        .footer {{
            background: #f8f9fa;
            padding: 20px;
            text-align: center;
            margin-top: 30px;
        }}
        .signature {{
            display: flex;
            justify-content: space-between;
            max-width: 600px;
            margin: 15px auto 0;
            font-size: 14px;
        }}
        .left-panel .personal-info {{
            background: rgba(255,255,255,0.1);
            padding: 15px;
            border-radius: 5px;
            margin-top: 15px;
        }}
        .left-panel .info-row {{
            margin-bottom: 8px;
            font-size: 13px;
        }}
        .left-panel .info-label {{
            font-weight: bold;
            color: {colors['accent']};
            display: inline-block;
            width: 90px;
        }}
        ul {{
            list-style-type: disc;
        }}
        @media print {{
            body {{
                background: white;
                padding: 0;
            }}
        }}
    </style>
</head>
<body>
    <div class="resume">
        <div class="left-panel">
            <div class="profile-image">
                {photo_html}
            </div>
            
            <div class="section-title">CONTACT</div>
            <div class="contact-item">📧 {data['personal']['email'] if data['personal']['email'] else 'email@example.com'}</div>
            <div class="contact-item">📱 {data['personal']['phone'] if data['personal']['phone'] else '+91 00000 00000'}</div>
            <div class="contact-item">📍 {data['personal']['city'] if data['personal']['city'] else 'City'}, {data['personal']['state'] if data['personal']['state'] else 'State'}</div>
            
            <div class="section-title">CORE EXPERTISE</div>
            <div style="font-size: 13px; line-height: 1.6;">{skills_text}</div>
            
            <div class="section-title">EDUCATION</div>
            {edu_html}
            
            <div class="section-title">CERTIFICATIONS</div>
            {cert_html}
            
            <div class="section-title">LANGUAGES</div>
            <div style="font-size: 13px;">{lang_text}</div>
            
            <div class="section-title">PERSONAL DETAILS</div>
            <div class="personal-info">
                <div class="info-row"><span class="info-label">Father:</span> {personal['father_name'] if personal.get('father_name') else 'Not specified'}</div>
                <div class="info-row"><span class="info-label">DOB:</span> {personal['dob'] if personal.get('dob') else 'Not specified'}</div>
                <div class="info-row"><span class="info-label">Marital:</span> {personal['marital_status'] if personal.get('marital_status') else 'Not specified'}</div>
                <div class="info-row"><span class="info-label">Nationality:</span> {personal['nationality'] if personal.get('nationality') else 'Not specified'}</div>
            </div>
        </div>
        
        <div class="right-panel">
            <h1 class="name">{data['personal']['name'] if data['personal']['name'] else 'Your Name'}</h1>
            <div class="title">{data['personal']['title'] if data['personal']['title'] else 'Professional Title'}</div>
            
            <div class="stats-grid">
                <div class="stat-item"><div class="stat-number">{len(data['experience'])}</div><div class="stat-label">Roles</div></div>
                <div class="stat-item"><div class="stat-number">{len(data['education'])}</div><div class="stat-label">Entries</div></div>
                <div class="stat-item"><div class="stat-number">{total_skills}</div><div class="stat-label">Skills</div></div>
                <div class="stat-item"><div class="stat-number">{len(data['certifications'])}</div><div class="stat-label">Certs</div></div>
            </div>
            
            <div class="section-title">PROFESSIONAL SUMMARY</div>
            <p style="font-size: 14px; line-height: 1.6;">{data['summary'] if data['summary'] else 'Add your professional summary in the editor.'}</p>
            
            <div class="section-title">WORK EXPERIENCE</div>
            {exp_html}
            
            <div class="section-title">KEY ACHIEVEMENTS</div>
            <ul>
                {ach_html}
            </ul>
            
            <div class="footer">
                <p>I hereby declare that the information provided is true and correct.</p>
                <div class="signature">
                    <span>Date: _____________</span>
                    <span>Place: {data['personal']['city'] if data['personal']['city'] else '______'}</span>
                    <span style="font-weight: bold; color: {colors['primary']};">({data['personal']['name'] if data['personal']['name'] else 'Your Name'})</span>
                </div>
            </div>
        </div>
    </div>
</body>
</html>"""
    return html

# ========== PDF-FRIENDLY STYLE 2: MODERN CORPORATE ==========
def generate_style2_html(template_id, data, template_colors=None):
    if template_colors:
        colors = template_colors
    else:
        all_templates = get_all_templates()
        template = all_templates.get(template_id)
        colors = template['colors']
    
    years_exp = 13
    total_skills = sum(len(skills) for skills in data['skills'].values()) if data['skills'] else 0
    
    # Photo handling
    photo_html = ""
    if data.get('profile_image'):
        photo_html = f'<img src="{data["profile_image"]}" style="width:100%; height:100%; border-radius:50%; object-fit:cover;" alt="Profile">'
    else:
        photo_html = f'<span style="font-size:40px; color:{colors["primary"]};">{data["personal"]["name"][0] if data["personal"]["name"] else "?"}</span>'
    
    # Build skills as badges
    skills_html = ""
    if data['skills']:
        for category, skills in data['skills'].items():
            for skill in skills:
                skills_html += f'<span style="display:inline-block; background:{colors["accent"]}; padding:5px 12px; border-radius:20px; margin:0 5px 10px 0; font-size:12px;">{skill}</span>'
    else:
        skills_html = '<span style="color: #888; font-style: italic;">Add skills in editor</span>'
    
    # Build education - Flexible fields
    edu_html = ""
    for edu in data['education']:
        edu_html += f"""
            <div style="background:white; padding:12px; margin-bottom:12px; border-left:4px solid {colors['secondary']};">"""
        
        if edu.get('year'):
            edu_html += f'<div style="color:{colors["secondary"]}; font-weight:bold;">{edu["year"]}</div>'
        if edu.get('degree'):
            edu_html += f'<div style="font-weight:bold;">{edu["degree"]}</div>'
        if edu.get('institution'):
            edu_html += f'<div style="color:#666;">{edu["institution"]}</div>'
        if edu.get('grade'):
            edu_html += f'<div style="color:{colors["secondary"]};">{edu["grade"]}</div>'
        if edu.get('specialization'):
            edu_html += f'<div style="color:#666; font-size:12px;">{edu["specialization"]}</div>'
        
        edu_html += "</div>"
    if not data['education']:
        edu_html = '<div style="color: #888; font-style: italic;">No education added.</div>'
    
    # Build certifications
    cert_html = ""
    for cert in data['certifications']:
        cert_html += f"""
            <div style="background:white; padding:12px; margin-bottom:12px; border-left:4px solid {colors['secondary']};">
                <div style="color:{colors['secondary']}; font-weight:bold;">{cert['year'] if cert.get('year') else 'Year'}</div>
                <div style="font-weight:bold;">{cert['name'] if cert.get('name') else 'Certification'}</div>
                <div style="color:#666;">{cert['issuer'] if cert.get('issuer') else 'Issuer'}</div>
            </div>
        """
    if not data['certifications']:
        cert_html = '<div style="color: #888; font-style: italic;">No certifications added.</div>'
    
    # Build languages
    lang_html = ""
    if data['languages']:
        for lang in data['languages']:
            lang_html += f'<span style="display:inline-block; background:{colors["accent"]}; padding:5px 15px; border-radius:20px; margin:0 5px 10px 0;">{lang["name"]} ({lang["proficiency"]})</span>'
    else:
        lang_html = '<span style="color: #888; font-style: italic;">Add languages</span>'
    
    # Build experience
    exp_html = ""
    for exp in data['experience']:
        exp_html += f"""
            <div style="margin-bottom:25px; padding-left:20px; position:relative;">
                <div style="position:absolute; left:0; top:8px; width:12px; height:12px; background:{colors['secondary']}; border-radius:50%;"></div>
                <div style="display:flex; justify-content:space-between; margin-bottom:5px; flex-wrap:wrap;">
                    <span style="font-weight:bold; color:{colors['primary']};">{exp['company'] if exp.get('company') else 'Company'}</span>
                    <span style="color:{colors['secondary']};">{exp['start_date'] if exp.get('start_date') else 'Start'}-{exp['end_date'] if exp.get('end_date') else 'End'}</span>
                </div>
                <div style="color:#666; margin:5px 0;">{exp['position'] if exp.get('position') else 'Position'} | {exp['location'] if exp.get('location') else 'Location'}</div>
                <ul style="margin-left:20px; margin-top:5px;">
        """
        if exp.get('description'):
            for desc in exp['description']:
                exp_html += f"<li style='margin-bottom:5px;'>{desc}</li>"
        else:
            exp_html += "<li style='color: #888;'>Add description</li>"
        exp_html += "</ul></div>"
    if not data['experience']:
        exp_html = '<div style="color: #888; font-style: italic;">No experience added.</div>'
    
    # Build achievements
    ach_html = ""
    for achievement in data['achievements']:
        ach_html += f'<span style="display:inline-block; background:{colors["accent"]}; padding:5px 15px; border-radius:25px; margin:0 5px 10px 0;">{achievement}</span>'
    if not data['achievements']:
        ach_html = '<span style="color: #888; font-style: italic;">No achievements added.</span>'
    
    personal = data['personal_details']
    
    html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{data['personal']['name'] if data['personal']['name'] else 'Your Name'} - Resume</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ background: {colors['primary']}; padding: 30px; font-family: Arial, sans-serif; }}
        .resume {{ max-width: 1000px; margin: 0 auto; background: white; border-radius: 10px; overflow: hidden; }}
        .header {{
            background: linear-gradient(90deg, {colors['primary']}, {colors['secondary']});
            padding: 30px;
            color: white;
        }}
        .name {{ font-size: 42px; font-weight: bold; margin-bottom: 5px; }}
        .contact {{ display: flex; gap: 20px; margin-top: 10px; flex-wrap: wrap; }}
        .profile-container {{ display: flex; align-items: center; gap: 20px; }}
        .profile-photo {{
            width: 80px;
            height: 80px;
            border-radius: 50%;
            background: white;
            display: flex;
            align-items: center;
            justify-content: center;
            overflow: hidden;
        }}
        .content {{ display: flex; }}
        .sidebar {{ width: 35%; background: #f5f8ff; padding: 25px; }}
        .main {{ width: 65%; padding: 25px; }}
        .section-title {{ font-size: 18px; font-weight: bold; color: {colors['primary']}; margin: 20px 0 10px; border-bottom: 3px solid {colors['secondary']}; padding-bottom: 5px; }}
        .stat-card {{ background: white; padding: 15px; margin-bottom: 15px; border-left: 5px solid {colors['secondary']}; }}
        .stat-number {{ font-size: 28px; font-weight: bold; color: {colors['primary']}; }}
        .footer {{ background: {colors['primary']}; color: white; padding: 20px; text-align: center; }}
        .signature {{ display: flex; justify-content: space-between; margin-top: 15px; flex-wrap: wrap; }}
    </style>
</head>
<body>
    <div class="resume">
        <div class="header">
            <div class="profile-container">
                <div class="profile-photo">
                    {photo_html}
                </div>
                <div>
                    <h1 class="name">{data['personal']['name'] if data['personal']['name'] else 'Your Name'}</h1>
                    <div>{data['personal']['title'] if data['personal']['title'] else 'Professional Title'}</div>
                </div>
            </div>
            <div class="contact">
                <span>📧 {data['personal']['email'] if data['personal']['email'] else 'email@example.com'}</span>
                <span>📱 {data['personal']['phone'] if data['personal']['phone'] else '+91 00000 00000'}</span>
                <span>📍 {data['personal']['city'] if data['personal']['city'] else 'City'}</span>
            </div>
        </div>
        
        <div class="content">
            <div class="sidebar">
                <div class="section-title">Key Metrics</div>
                <div class="stat-card"><div class="stat-number">{len(data['experience'])}</div><div>Roles</div></div>
                <div class="stat-card"><div class="stat-number">{len(data['education'])}</div><div>Entries</div></div>
                
                <div class="section-title">Skills</div>
                <div>{skills_html}</div>
                
                <div class="section-title">Education</div>
                {edu_html}
                
                <div class="section-title">Certifications</div>
                {cert_html}
                
                <div class="section-title">Languages</div>
                <div>{lang_html}</div>
                
                <div class="section-title">Personal</div>
                <div style="background:white; padding:12px; border-radius:8px;">
                    <div><strong>Father:</strong> {personal['father_name'] if personal.get('father_name') else 'Not specified'}</div>
                    <div><strong>DOB:</strong> {personal['dob'] if personal.get('dob') else 'Not specified'}</div>
                    <div><strong>Marital:</strong> {personal['marital_status'] if personal.get('marital_status') else 'Not specified'}</div>
                    <div><strong>Nationality:</strong> {personal['nationality'] if personal.get('nationality') else 'Not specified'}</div>
                </div>
            </div>
            
            <div class="main">
                <div class="section-title">Summary</div>
                <p>{data['summary'] if data['summary'] else 'Add your professional summary in the editor.'}</p>
                
                <div class="section-title">Experience</div>
                {exp_html}
                
                <div class="section-title">Achievements</div>
                <div>{ach_html}</div>
                
                <div class="footer">
                    <p>I hereby declare that the information is true and correct.</p>
                    <div class="signature">
                        <span>Date: _____________</span>
                        <span>Place: {data['personal']['city'] if data['personal']['city'] else '______'}</span>
                        <span style="font-weight:bold;">({data['personal']['name'] if data['personal']['name'] else 'Your Name'})</span>
                    </div>
                </div>
            </div>
        </div>
    </div>
</body>
</html>"""
    return html

# ========== PDF-FRIENDLY STYLE 3: CREATIVE HOSPITALITY ==========
def generate_style3_html(template_id, data, template_colors=None):
    if template_colors:
        colors = template_colors
    else:
        all_templates = get_all_templates()
        template = all_templates.get(template_id)
        colors = template['colors']
    
    # Photo handling
    photo_html = ""
    if data.get('profile_image'):
        photo_html = f'<img src="{data["profile_image"]}" style="width:80px; height:80px; border-radius:50%; object-fit:cover;" alt="Profile">'
    else:
        photo_html = f'<div style="width:80px; height:80px; border-radius:50%; background:{colors["primary"]}; display:flex; align-items:center; justify-content:center; font-size:40px; color:white;">{data["personal"]["name"][0] if data["personal"]["name"] else "?"}</div>'
    
    # Build skills as pills
    skills_html = ""
    if data['skills']:
        for category, skills in data['skills'].items():
            for skill in skills:
                skills_html += f'<span style="display:inline-block; background:white; padding:5px 15px; border-radius:30px; margin:0 5px 10px 0; border:1px solid {colors["primary"]};">{skill}</span>'
    else:
        skills_html = '<span style="color: #888; font-style: italic;">Add skills</span>'
    
    # Build education - Flexible fields
    edu_html = ""
    for edu in data['education']:
        edu_html += f"""
            <div style="margin-bottom:15px;">"""
        
        if edu.get('year'):
            edu_html += f'<div style="font-size:18px; font-weight:bold; color:{colors["primary"]};">{edu["year"]}</div>'
        if edu.get('degree'):
            edu_html += f'<div style="font-weight:bold;">{edu["degree"]}</div>'
        if edu.get('institution'):
            edu_html += f'<div style="color:#666;">{edu["institution"]}</div>'
        if edu.get('grade'):
            edu_html += f'<div style="color:{colors["primary"]};">{edu["grade"]}</div>'
        if edu.get('specialization'):
            edu_html += f'<div style="color:#666; font-size:13px;">{edu["specialization"]}</div>'
        
        edu_html += "</div>"
    if not data['education']:
        edu_html = '<div style="color: #888; font-style: italic;">No education added.</div>'
    
    # Build certifications
    cert_html = ""
    for cert in data['certifications']:
        cert_html += f"""
            <div style="margin-bottom:12px;">
                <div style="color:{colors['primary']}; font-weight:bold;">{cert['year'] if cert.get('year') else 'Year'}</div>
                <div style="font-weight:bold;">{cert['name'] if cert.get('name') else 'Certification'}</div>
                <div style="color:#666;">{cert['issuer'] if cert.get('issuer') else 'Issuer'}</div>
            </div>
        """
    if not data['certifications']:
        cert_html = '<div style="color: #888; font-style: italic;">No certifications added.</div>'
    
    # Build languages
    lang_html = ""
    if data['languages']:
        for lang in data['languages']:
            lang_html += f'<span style="display:inline-block; background:white; padding:5px 15px; border-radius:30px; margin:0 5px 10px 0; border:1px solid {colors["primary"]};"><strong>{lang["name"]}</strong> ({lang["proficiency"]})</span>'
    else:
        lang_html = '<span style="color: #888; font-style: italic;">Add languages</span>'
    
    # Build experience
    exp_html = ""
    for exp in data['experience']:
        exp_html += f"""
            <div style="background:white; padding:20px; border-radius:15px; margin-bottom:20px; border:1px solid {colors['primary']};">
                <div style="display:flex; justify-content:space-between; margin-bottom:10px; flex-wrap:wrap;">
                    <span style="font-size:20px; font-weight:bold; color:{colors['secondary']};">{exp['company'] if exp.get('company') else 'Company'}</span>
                    <span style="color:{colors['primary']};">{exp['start_date'] if exp.get('start_date') else 'Start'}-{exp['end_date'] if exp.get('end_date') else 'End'}</span>
                </div>
                <div style="color:#666; margin:5px 0;">{exp['position'] if exp.get('position') else 'Position'} | {exp['location'] if exp.get('location') else 'Location'}</div>
                <ul style="margin-left:20px; margin-top:10px;">
        """
        if exp.get('description'):
            for desc in exp['description']:
                exp_html += f"<li style='margin-bottom:5px;'>{desc}</li>"
        else:
            exp_html += "<li style='color: #888;'>Add description</li>"
        exp_html += "</ul></div>"
    if not data['experience']:
        exp_html = '<div style="color: #888; font-style: italic;">No experience added.</div>'
    
    # Build achievements
    ach_html = ""
    for achievement in data['achievements']:
        ach_html += f"<li style='margin-bottom:5px;'>{achievement}</li>"
    if not data['achievements']:
        ach_html = '<li style="color: #888; font-style: italic;">No achievements added.</li>'
    
    personal = data['personal_details']
    
    html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{data['personal']['name'] if data['personal']['name'] else 'Your Name'} - Resume</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ background: {colors['secondary']}; padding: 30px; font-family: Georgia, serif; }}
        .resume {{ max-width: 1000px; margin: 0 auto; background: #fff9f5; border-radius: 20px; overflow: hidden; }}
        .header {{
            background: linear-gradient(135deg, {colors['primary']}, {colors['secondary']});
            padding: 40px;
            text-align: center;
            color: white;
        }}
        .name {{ font-size: 54px; font-weight: bold; margin-bottom: 10px; }}
        .profile-container {{ display: flex; justify-content: center; margin-bottom: 20px; }}
        .contact {{
            display: flex;
            justify-content: center;
            gap: 30px;
            background: white;
            padding: 15px 30px;
            border-radius: 50px;
            width: fit-content;
            margin: -25px auto 0;
            border: 2px solid {colors['primary']};
            flex-wrap: wrap;
        }}
        .content {{
            padding: 40px;
            display: grid;
            grid-template-columns: 1fr 2fr;
            gap: 30px;
        }}
        .left {{
            background: {colors['accent']};
            padding: 25px;
            border-radius: 20px;
        }}
        .section-title {{
            font-size: 22px;
            font-weight: bold;
            color: {colors['secondary']};
            margin: 20px 0 15px;
            border-bottom: 2px dashed {colors['primary']};
            padding-bottom: 5px;
        }}
        .left .section-title {{
            color: {colors['secondary']};
        }}
        .footer {{
            background: linear-gradient(145deg, {colors['primary']}, {colors['secondary']});
            padding: 25px;
            color: white;
            text-align: center;
            margin-top: 20px;
        }}
        .signature {{ display: flex; justify-content: space-between; margin-top: 15px; flex-wrap: wrap; }}
        ul {{
            list-style-type: disc;
        }}
    </style>
</head>
<body>
    <div class="resume">
        <div class="header">
            <div class="profile-container">
                {photo_html}
            </div>
            <h1 class="name">{data['personal']['name'] if data['personal']['name'] else 'Your Name'}</h1>
            <div>{data['personal']['title'] if data['personal']['title'] else 'Professional Title'}</div>
        </div>
        
        <div class="contact">
            <span>📧 {data['personal']['email'] if data['personal']['email'] else 'email@example.com'}</span>
            <span>📱 {data['personal']['phone'] if data['personal']['phone'] else '+91 00000 00000'}</span>
            <span>📍 {data['personal']['city'] if data['personal']['city'] else 'City'}</span>
        </div>
        
        <div class="content">
            <div class="left">
                <div class="section-title">Skills</div>
                <div>{skills_html}</div>
                
                <div class="section-title">Education</div>
                {edu_html}
                
                <div class="section-title">Certifications</div>
                {cert_html}
                
                <div class="section-title">Languages</div>
                <div>{lang_html}</div>
                
                <div class="section-title">Personal</div>
                <div style="background:white; padding:15px; border-radius:10px;">
                    <div><strong>Father:</strong> {personal['father_name'] if personal.get('father_name') else 'Not specified'}</div>
                    <div><strong>DOB:</strong> {personal['dob'] if personal.get('dob') else 'Not specified'}</div>
                    <div><strong>Marital:</strong> {personal['marital_status'] if personal.get('marital_status') else 'Not specified'}</div>
                    <div><strong>Nationality:</strong> {personal['nationality'] if personal.get('nationality') else 'Not specified'}</div>
                </div>
            </div>
            
            <div class="right">
                <div style="background:white; padding:25px; border-radius:15px; margin-bottom:20px; border:1px solid {colors['primary']};">
                    <p style="font-size:16px; line-height:1.6;">{data['summary'] if data['summary'] else 'Add your professional summary in the editor.'}</p>
                </div>
                
                <div class="section-title">Experience</div>
                {exp_html}
                
                <div class="section-title">Achievements</div>
                <ul>
                    {ach_html}
                </ul>
                
                <div class="footer">
                    <p>I hereby declare that the information is true and correct.</p>
                    <div class="signature">
                        <span>Date: _____________</span>
                        <span>Place: {data['personal']['city'] if data['personal']['city'] else '______'}</span>
                        <span style="font-weight:bold;">({data['personal']['name'] if data['personal']['name'] else 'Your Name'})</span>
                    </div>
                </div>
            </div>
        </div>
    </div>
</body>
</html>"""
    return html

# Generate PDF using pdfkit
def generate_pdf(html_content):
    """Convert HTML to PDF using pdfkit"""
    try:
        options = {
            'page-size': 'A4',
            'margin-top': '0.5in',
            'margin-right': '0.5in',
            'margin-bottom': '0.5in',
            'margin-left': '0.5in',
            'encoding': "UTF-8",
            'no-outline': None,
            'enable-local-file-access': None
        }
        
        with tempfile.NamedTemporaryFile(suffix='.html', delete=False, mode='w', encoding='utf-8') as f:
            f.write(html_content)
            html_path = f.name
        
        pdf = pdfkit.from_file(html_path, False, options=options)
        os.unlink(html_path)
        return pdf
    except Exception as e:
        st.error(f"PDF Error: {str(e)}")
        st.info("💡 Please install wkhtmltopdf")
        return None

# Generate Word Document with theme colors
def generate_word_doc(data, template_id):
    all_templates = get_all_templates()
    template = all_templates.get(template_id)
    colors = template['colors']
    
    def hex_to_rgb(hex_color):
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    
    primary_rgb = hex_to_rgb(colors['primary'])
    secondary_rgb = hex_to_rgb(colors['secondary'])
    
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = header.add_run(data['personal']['name'] if data['personal']['name'] else 'Your Name')
    name_run.font.size = Pt(36)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(*primary_rgb)
    header.add_run('\n')
    
    title_run = header.add_run(data['personal']['title'] if data['personal']['title'] else 'Professional Title')
    title_run.font.size = Pt(16)
    title_run.font.italic = True
    title_run.font.color.rgb = RGBColor(*secondary_rgb)
    header.add_run('\n\n')
    
    # Contact
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(f"📧 {data['personal']['email'] if data['personal']['email'] else 'email@example.com'}     📱 {data['personal']['phone'] if data['personal']['phone'] else '+91 00000 00000'}     📍 {data['personal']['city'] if data['personal']['city'] else 'City'}")
    
    doc.add_paragraph()
    
    # Summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(data['summary'] if data['summary'] else 'Add your professional summary here.')
    doc.add_paragraph()
    
    # Experience
    doc.add_heading('Work Experience', level=1)
    if data['experience']:
        for exp in data['experience']:
            p = doc.add_paragraph()
            p.add_run(f"{exp['company'] if exp.get('company') else 'Company'} - {exp['location'] if exp.get('location') else 'Location'}").bold = True
            p.add_run(f"\n{exp['position'] if exp.get('position') else 'Position'}").italic = True
            p.add_run(f"    ({exp['start_date'] if exp.get('start_date') else 'Start'} - {exp['end_date'] if exp.get('end_date') else 'End'})")
            if exp.get('description'):
                for desc in exp['description']:
                    doc.add_paragraph(desc, style='List Bullet')
            doc.add_paragraph()
    else:
        doc.add_paragraph("No experience added yet. Click 'Edit All Sections' to add.")
        doc.add_paragraph()
    
    # Education - Flexible fields
    doc.add_heading('Education', level=1)
    if data['education']:
        for edu in data['education']:
            line = []
            if edu.get('degree'):
                line.append(edu['degree'])
            if edu.get('institution'):
                line.append(f"from {edu['institution']}")
            if edu.get('year'):
                line.append(f"({edu['year']})")
            
            if line:
                p = doc.add_paragraph()
                p.add_run(" ".join(line)).bold = True
            
            if edu.get('grade'):
                doc.add_paragraph(f"Grade: {edu['grade']}", style='List Bullet')
            if edu.get('specialization'):
                doc.add_paragraph(f"Specialization: {edu['specialization']}", style='List Bullet')
        doc.add_paragraph()
    else:
        doc.add_paragraph("No education added yet. Click 'Edit All Sections' to add.")
        doc.add_paragraph()
    
    # Skills
    doc.add_heading('Skills & Expertise', level=1)
    if data['skills']:
        for category, skills in data['skills'].items():
            doc.add_paragraph(f"{category}: {', '.join(skills)}")
        doc.add_paragraph()
    else:
        doc.add_paragraph("No skills added yet. Click 'Edit All Sections' to add.")
        doc.add_paragraph()
    
    # Certifications
    if data['certifications']:
        doc.add_heading('Certifications', level=1)
        for cert in data['certifications']:
            doc.add_paragraph(f"• {cert['name'] if cert.get('name') else 'Certification'} - {cert['issuer'] if cert.get('issuer') else 'Issuer'} ({cert['year'] if cert.get('year') else 'Year'})")
        doc.add_paragraph()
    
    # Languages
    if data['languages']:
        doc.add_heading('Languages', level=1)
        for lang in data['languages']:
            doc.add_paragraph(f"• {lang['name'] if lang.get('name') else 'Language'} ({lang['proficiency'] if lang.get('proficiency') else 'Proficiency'})")
        doc.add_paragraph()
    
    # Achievements
    doc.add_heading('Key Achievements', level=1)
    if data['achievements']:
        for achievement in data['achievements']:
            doc.add_paragraph(f"• {achievement}")
        doc.add_paragraph()
    else:
        doc.add_paragraph("No achievements added yet. Click 'Edit All Sections' to add.")
        doc.add_paragraph()
    
    # Personal Details
    doc.add_heading('Personal Details', level=1)
    personal = data['personal_details']
    doc.add_paragraph(f"Father's Name: {personal['father_name'] if personal.get('father_name') else 'Not specified'}")
    doc.add_paragraph(f"Date of Birth: {personal['dob'] if personal.get('dob') else 'Not specified'}")
    doc.add_paragraph(f"Marital Status: {personal['marital_status'] if personal.get('marital_status') else 'Not specified'}")
    doc.add_paragraph(f"Nationality: {personal['nationality'] if personal.get('nationality') else 'Not specified'}")
    doc.add_paragraph()
    
    # Declaration
    decl = doc.add_paragraph()
    decl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    decl.add_run("I hereby declare that the information provided is true and correct.").italic = True
    
    doc.add_paragraph()
    
    # Signature
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig.add_run(f"Date: _____________    Place: {data['personal']['city'] if data['personal']['city'] else '______'}    ({data['personal']['name'] if data['personal']['name'] else 'Your Name'})")
    
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes.getvalue()

def main():
    # Initialize session state if not already done
    if 'resume_data' not in st.session_state:
        init_session_state()
    
    st.markdown("""
    <div class="main-header">
        <h1>🎯 Professional Resume Generator - 1000+ Templates</h1>
        <p>✅ 105 Built-in • Create Unlimited Custom Templates • Photo Upload • PDF & Word Download • All Sections</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar
    with st.sidebar:
        st.image("https://img.icons8.com/fluency/96/resume.png", width=80)
        st.title("Resume Builder")
        st.markdown("---")
        
        # Photo Upload
        with st.expander("📸 Profile Photo", expanded=True):
            uploaded_file = st.file_uploader("Upload Photo", type=['jpg', 'jpeg', 'png'])
            if uploaded_file:
                bytes_data = uploaded_file.getvalue()
                encoded = base64.b64encode(bytes_data).decode()
                if st.session_state.resume_data is not None:
                    st.session_state.resume_data['profile_image'] = f"data:image/{uploaded_file.type.split('/')[-1]};base64,{encoded}"
                st.image(uploaded_file, width=100)
                st.success("✅ Photo uploaded!")
        
        # Template Family Selection
        st.subheader("🎨 Template Family")
        
        # Get all templates
        all_templates = get_all_templates()
        
        # Create filter options including custom templates
        families = ['All'] + sorted(list(set([t['style_family'] for t in all_templates.values()])))
        
        selected_filter = st.selectbox("Filter by Style", families)
        st.session_state.template_filter = selected_filter
        
        # Add New Template Section
        with st.expander("➕ Add New Custom Template", expanded=False):
            st.markdown("### Create Your Own Template")
            template_name = st.text_input("Template Name", "My Custom Template")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                primary_color = st.color_picker("Primary Color", "#1e2b3a")
            with col2:
                secondary_color = st.color_picker("Secondary Color", "#ff6b6b")
            with col3:
                accent_color = st.color_picker("Accent Color", "#ffd93d")
            
            style_type = st.selectbox("Base Style", 
                ['Minimalist Luxury', 'Modern Corporate', 'Creative Hospitality', 'Custom Mix'])
            
            if st.button("✨ Create Template", use_container_width=True):
                new_id = add_custom_template(template_name, primary_color, secondary_color, accent_color, style_type)
                st.success(f"✅ Template '{template_name}' created!")
                st.session_state.selected_template = new_id
                st.rerun()
        
        st.markdown(f"**Total Templates:** {len(all_templates)}")
        st.markdown(f"**Built-in:** 105 | **Custom:** {len(st.session_state.custom_templates) if 'custom_templates' in st.session_state else 0}")
        st.markdown("---")
        
        # Personal Information
        with st.expander("👤 Personal Info", expanded=True):
            if st.session_state.resume_data is not None:
                current_name = st.session_state.resume_data.get('personal', {}).get('name', '')
                current_email = st.session_state.resume_data.get('personal', {}).get('email', '')
                current_phone = st.session_state.resume_data.get('personal', {}).get('phone', '')
                current_city = st.session_state.resume_data.get('personal', {}).get('city', '')
                current_state = st.session_state.resume_data.get('personal', {}).get('state', '')
                current_title = st.session_state.resume_data.get('personal', {}).get('title', '')
            else:
                current_name = current_email = current_phone = current_city = current_state = current_title = ''
            
            name = st.text_input("Full Name", current_name)
            email = st.text_input("Email", current_email)
            phone = st.text_input("Phone", current_phone)
            title = st.text_input("Professional Title", current_title)
            city = st.text_input("City", current_city)
            state = st.text_input("State", current_state)
            
            if st.session_state.resume_data is not None:
                st.session_state.resume_data['personal']['name'] = name
                st.session_state.resume_data['personal']['email'] = email
                st.session_state.resume_data['personal']['phone'] = phone
                st.session_state.resume_data['personal']['title'] = title
                st.session_state.resume_data['personal']['city'] = city
                st.session_state.resume_data['personal']['state'] = state
        
        # Summary
        with st.expander("📝 Summary", expanded=True):
            if st.session_state.resume_data is not None:
                current_summary = st.session_state.resume_data.get('summary', '')
            else:
                current_summary = ''
            summary = st.text_area("Professional Summary", current_summary, height=100)
            if st.session_state.resume_data is not None:
                st.session_state.resume_data['summary'] = summary
        
        # Reset Button
        if st.button("🔄 Reset to Blank", use_container_width=True):
            safe_reset()
            st.rerun()
    
    # Main content
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.markdown("### 🎨 Templates")
        
        # Filter templates
        all_templates = get_all_templates()
        if st.session_state.template_filter == 'All':
            filtered_templates = all_templates
        else:
            filtered_templates = {tid: t for tid, t in all_templates.items() 
                                 if t['style_family'] == st.session_state.template_filter}
        
        # Show custom templates first with badge
        custom_templates = {tid: t for tid, t in filtered_templates.items() if t.get('is_custom', False)}
        builtin_templates = {tid: t for tid, t in filtered_templates.items() if not t.get('is_custom', False)}
        
        # Display custom templates first
        all_display = list(custom_templates.items()) + list(builtin_templates.items())
        
        for i, (tid, template) in enumerate(all_display[:12]):
            is_selected = st.session_state.selected_template == tid
            card_class = "template-card selected" if is_selected else "template-card"
            
            # Add badge for custom templates
            badge = "✨ " if template.get('is_custom', False) else ""
            
            st.markdown(f"""
            <div class="{card_class}">
                <div class="template-preview" style="background: linear-gradient(135deg, {template['colors']['primary']} 0%, {template['colors']['secondary']} 100%);">
                    <span>{badge}📄</span>
                </div>
                <h4>{badge}{template['name']}</h4>
                <div style="display: flex; gap: 3px; margin-top: 5px;">
                    <div style="width: 20px; height: 20px; background: {template['colors']['primary']}; border-radius: 5px;"></div>
                    <div style="width: 20px; height: 20px; background: {template['colors']['secondary']}; border-radius: 5px;"></div>
                    <div style="width: 20px; height: 20px; background: {template['colors']['accent']}; border-radius: 5px;"></div>
                </div>
                <p style="font-size: 10px; color: #666;">{template['style_family']}</p>
            </div>
            """, unsafe_allow_html=True)
            
            col_btn1, col_btn2 = st.columns([3, 1])
            with col_btn1:
                if st.button("Select", key=f"sel_{tid}"):
                    st.session_state.selected_template = tid
                    st.rerun()
            with col_btn2:
                if template.get('is_custom', False):
                    if st.button("🗑️", key=f"del_{tid}"):
                        if tid in st.session_state.custom_templates:
                            del st.session_state.custom_templates[tid]
                            if st.session_state.selected_template == tid:
                                st.session_state.selected_template = 'style1_v1'
                            st.rerun()
    
    with col2:
        st.markdown("### 👁️ Preview")
        
        if st.session_state.resume_data is not None:
            selected_tid = st.session_state.selected_template
            all_templates = get_all_templates()
            
            if selected_tid in all_templates:
                template = all_templates[selected_tid]
                
                # Determine which style to use based on template's style family
                style_family = template['style_family']
                
                if 'Minimalist Luxury' in style_family:
                    html = generate_style1_html(selected_tid, st.session_state.resume_data, template['colors'])
                elif 'Modern Corporate' in style_family:
                    html = generate_style2_html(selected_tid, st.session_state.resume_data, template['colors'])
                elif 'Creative Hospitality' in style_family:
                    html = generate_style3_html(selected_tid, st.session_state.resume_data, template['colors'])
                else:
                    # Default to style 1 for custom mix
                    html = generate_style1_html(selected_tid, st.session_state.resume_data, template['colors'])
                
                st.components.v1.html(html, height=700, scrolling=True)
                
                # Download
                st.markdown("### 📥 Download")
                col_a, col_b, col_c = st.columns(3)
                
                with col_a:
                    b64 = base64.b64encode(html.encode()).decode()
                    st.markdown(f'<a href="data:text/html;base64,{b64}" download="resume.html"><button class="download-btn">📄 HTML</button></a>', unsafe_allow_html=True)
                
                with col_b:
                    if st.button("📑 Generate PDF", use_container_width=True):
                        with st.spinner("Generating PDF..."):
                            pdf_bytes = generate_pdf(html)
                            if pdf_bytes:
                                st.session_state.pdf_data = base64.b64encode(pdf_bytes).decode()
                                st.success("✅ PDF Ready!")
                    
                    if st.session_state.pdf_data:
                        st.markdown(f'<a href="data:application/pdf;base64,{st.session_state.pdf_data}" download="resume.pdf" class="download-link">📑 Download PDF</a>', unsafe_allow_html=True)
                
                with col_c:
                    if st.button("📝 Generate WORD", use_container_width=True):
                        with st.spinner("Generating Word..."):
                            word_bytes = generate_word_doc(st.session_state.resume_data, selected_tid)
                            if word_bytes:
                                st.session_state.word_data = base64.b64encode(word_bytes).decode()
                                st.success("✅ Word Ready!")
                    
                    if st.session_state.word_data:
                        st.markdown(f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{st.session_state.word_data}" download="resume.docx" class="download-link">📝 Download WORD</a>', unsafe_allow_html=True)
            else:
                st.error("Template not found. Please select another template.")
    
    # Edit Sections
    with st.expander("✏️ Edit All Sections", expanded=False):
        if st.session_state.resume_data is not None:
            tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(["🎓 Education", "💼 Experience", "📜 Certifications", "🏆 Achievements", "🗣️ Languages", "📋 Personal Details"])
            
            with tab1:
                st.markdown("### Education")
                st.markdown("*Fields are flexible - you can enter any type of qualification*")
                
                # Option to add custom fields
                col1, col2 = st.columns([3, 1])
                with col1:
                    if st.button("➕ Add Education", key="add_edu"):
                        if 'education' not in st.session_state.resume_data:
                            st.session_state.resume_data['education'] = []
                        st.session_state.resume_data['education'].append({
                            'degree': '',
                            'institution': '',
                            'year': '',
                            'grade': '',
                            'specialization': ''
                        })
                        st.rerun()
                
                if 'education' in st.session_state.resume_data and st.session_state.resume_data['education']:
                    for i, edu in enumerate(st.session_state.resume_data['education']):
                        with st.container():
                            st.markdown(f"**Education #{i+1}**")
                            
                            # Flexible fields - user can enter anything
                            col1, col2 = st.columns(2)
                            with col1:
                                edu['degree'] = st.text_input("Degree/Certificate/Qualification", edu.get('degree', ''), key=f"edu_deg_{i}")
                                edu['institution'] = st.text_input("Institution/School", edu.get('institution', ''), key=f"edu_inst_{i}")
                            with col2:
                                edu['year'] = st.text_input("Year/Period", edu.get('year', ''), key=f"edu_year_{i}")
                                edu['grade'] = st.text_input("Grade/Score (optional)", edu.get('grade', ''), key=f"edu_grade_{i}")
                            
                            edu['specialization'] = st.text_input("Specialization/Field (optional)", edu.get('specialization', ''), key=f"edu_spec_{i}")
                            
                            col_del1, col_del2 = st.columns([5, 1])
                            with col_del2:
                                if st.button("🗑️ Delete", key=f"del_edu_{i}"):
                                    st.session_state.resume_data['education'].pop(i)
                                    st.rerun()
                            st.markdown("---")
                else:
                    st.info("No education entries. Click 'Add Education' to create one. You can enter any type of qualification (degree, diploma, certificate, etc.)")
            
            with tab2:
                st.markdown("### Work Experience")
                if st.button("➕ Add Experience", key="add_exp"):
                    if 'experience' not in st.session_state.resume_data:
                        st.session_state.resume_data['experience'] = []
                    st.session_state.resume_data['experience'].append({
                        'company': '',
                        'location': '',
                        'position': '',
                        'start_date': '',
                        'end_date': '',
                        'description': ['']
                    })
                    st.rerun()
                
                if 'experience' in st.session_state.resume_data and st.session_state.resume_data['experience']:
                    for i, exp in enumerate(st.session_state.resume_data['experience']):
                        with st.container():
                            st.markdown(f"**Experience #{i+1}**")
                            cols = st.columns([2, 2, 1, 1])
                            with cols[0]:
                                exp['company'] = st.text_input("Company", exp.get('company', ''), key=f"exp_comp_{i}")
                            with cols[1]:
                                exp['position'] = st.text_input("Position", exp.get('position', ''), key=f"exp_pos_{i}")
                            with cols[2]:
                                exp['start_date'] = st.text_input("Start", exp.get('start_date', ''), key=f"exp_start_{i}")
                            with cols[3]:
                                if st.button("🗑️", key=f"del_exp_{i}"):
                                    st.session_state.resume_data['experience'].pop(i)
                                    st.rerun()
                            
                            exp['location'] = st.text_input("Location", exp.get('location', ''), key=f"exp_loc_{i}")
                            exp['end_date'] = st.text_input("End Date", exp.get('end_date', ''), key=f"exp_end_{i}")
                            
                            desc_text = "\n".join(exp['description']) if exp['description'] else ""
                            new_desc = st.text_area("Description (one per line)", desc_text, key=f"exp_desc_{i}", height=80)
                            exp['description'] = [d.strip() for d in new_desc.split('\n') if d.strip()]
                            st.markdown("---")
                else:
                    st.info("No experience entries. Click 'Add Experience' to create one.")
            
            with tab3:
                st.markdown("### Certifications")
                if st.button("➕ Add Certification", key="add_cert"):
                    if 'certifications' not in st.session_state.resume_data:
                        st.session_state.resume_data['certifications'] = []
                    st.session_state.resume_data['certifications'].append({
                        'name': '',
                        'issuer': '',
                        'year': ''
                    })
                    st.rerun()
                
                if 'certifications' in st.session_state.resume_data and st.session_state.resume_data['certifications']:
                    for i, cert in enumerate(st.session_state.resume_data['certifications']):
                        cols = st.columns([2, 2, 1, 0.5])
                        with cols[0]:
                            cert['name'] = st.text_input("Name", cert.get('name', ''), key=f"cert_name_{i}")
                        with cols[1]:
                            cert['issuer'] = st.text_input("Issuer", cert.get('issuer', ''), key=f"cert_issuer_{i}")
                        with cols[2]:
                            cert['year'] = st.text_input("Year", cert.get('year', ''), key=f"cert_year_{i}")
                        with cols[3]:
                            if st.button("🗑️", key=f"del_cert_{i}"):
                                st.session_state.resume_data['certifications'].pop(i)
                                st.rerun()
                else:
                    st.info("No certifications. Click 'Add Certification' to create one.")
            
            with tab4:
                st.markdown("### Key Achievements")
                if 'achievements' not in st.session_state.resume_data:
                    st.session_state.resume_data['achievements'] = []
                
                achievements_text = "\n".join(st.session_state.resume_data['achievements'])
                new_achievements = st.text_area("Achievements (one per line)", achievements_text, height=150)
                st.session_state.resume_data['achievements'] = [a.strip() for a in new_achievements.split('\n') if a.strip()]
                
                if st.button("➕ Add Achievement", key="add_ach"):
                    st.session_state.resume_data['achievements'].append('')
                    st.rerun()
            
            with tab5:
                st.markdown("### Languages")
                if st.button("➕ Add Language", key="add_lang"):
                    if 'languages' not in st.session_state.resume_data:
                        st.session_state.resume_data['languages'] = []
                    st.session_state.resume_data['languages'].append({
                        'name': '',
                        'proficiency': 'Fluent'
                    })
                    st.rerun()
                
                if 'languages' in st.session_state.resume_data and st.session_state.resume_data['languages']:
                    for i, lang in enumerate(st.session_state.resume_data['languages']):
                        cols = st.columns([2, 2, 0.5])
                        with cols[0]:
                            lang['name'] = st.text_input("Language", lang.get('name', ''), key=f"lang_name_{i}")
                        with cols[1]:
                            lang['proficiency'] = st.selectbox("Proficiency", 
                                ['Native', 'Fluent', 'Professional', 'Intermediate', 'Basic'],
                                index=['Native', 'Fluent', 'Professional', 'Intermediate', 'Basic'].index(lang['proficiency']) 
                                if lang['proficiency'] in ['Native', 'Fluent', 'Professional', 'Intermediate', 'Basic'] else 1,
                                key=f"lang_prof_{i}")
                        with cols[2]:
                            if st.button("🗑️", key=f"del_lang_{i}"):
                                st.session_state.resume_data['languages'].pop(i)
                                st.rerun()
                else:
                    st.info("No languages. Click 'Add Language' to create one.")
            
            with tab6:
                st.markdown("### Personal Details")
                if 'personal_details' not in st.session_state.resume_data:
                    st.session_state.resume_data['personal_details'] = {}
                
                col1, col2 = st.columns(2)
                with col1:
                    father_name = st.text_input("Father's Name", st.session_state.resume_data['personal_details'].get('father_name', ''))
                    dob = st.text_input("Date of Birth", st.session_state.resume_data['personal_details'].get('dob', ''))
                with col2:
                    marital_status = st.selectbox("Marital Status", 
                        ['', 'Married', 'Unmarried', 'Divorced', 'Widowed'],
                        index=0 if not st.session_state.resume_data['personal_details'].get('marital_status') else 
                        ['', 'Married', 'Unmarried', 'Divorced', 'Widowed'].index(st.session_state.resume_data['personal_details']['marital_status']))
                    nationality = st.text_input("Nationality", st.session_state.resume_data['personal_details'].get('nationality', ''))
                
                st.session_state.resume_data['personal_details']['father_name'] = father_name
                st.session_state.resume_data['personal_details']['dob'] = dob
                st.session_state.resume_data['personal_details']['marital_status'] = marital_status
                st.session_state.resume_data['personal_details']['nationality'] = nationality
    
    st.markdown("""
    <div class="footer">
        <p>✅ 105 Built-in Templates • Create Unlimited Custom Templates • Photo Upload • PDF & Word Download • All Sections</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
